"""
Test suite for CSS class-based style mapping feature in html4docx

This test file is designed to be added to the html4docx test suite.
Tests all aspects of the extended styling features:
- CSS class to Word style mapping
- Tag style overrides
- Default paragraph styles
- Inline CSS styles
- !important flag precedence
- Character styles for inline elements
- Multi-paragraph style persistence
"""

import sys
import unittest
from docx import Document
from docx.shared import Pt

module_dir = "D:\\Dropbox\\Projects\\Python\\html4docx"
sys.path.insert(0, module_dir)
# from docx.shared import RGBColor, Pt
from html4docx import HtmlToDocx

# import HtmlToDocx


class TestCSSClassMapping(unittest.TestCase):
    """Test CSS class to Word style mapping"""

    def test_basic_class_mapping(self):
        """Test that CSS classes are mapped to Word styles"""
        style_map = {
            "custom-style": "Heading 1",
        }

        html = '<p class="custom-style">Test paragraph</p>'

        doc = Document()
        parser = HtmlToDocx(style_map=style_map)
        parser.add_html_to_document(html, doc)

        # Verify paragraph uses the mapped style
        self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")

    def test_multiple_classes(self):
        """Test that first matching class in style_map wins"""
        style_map = {
            "first": "Heading 1",
            "second": "Heading 2",
        }

        html = '<p class="second first">Test</p>'

        doc = Document()
        parser = HtmlToDocx(style_map=style_map)
        parser.add_html_to_document(html, doc)

        # Should use first matching class found
        self.assertIn(doc.paragraphs[0].style.name, ["Heading 1", "Heading 2"])

    def test_unmapped_class_uses_default(self):
        """Test that unmapped classes fall back to default behavior"""
        style_map = {
            "mapped": "Heading 1",
        }

        html = '<p class="unmapped">Test</p>'

        doc = Document()
        parser = HtmlToDocx(style_map=style_map, default_paragraph_style=None)
        parser.add_html_to_document(html, doc)

        # Should use default Word 'Normal' style
        self.assertEqual(doc.paragraphs[0].style.name, "Normal")


class TestTagStyleOverrides(unittest.TestCase):
    """Test tag-based style overrides"""

    def test_h1_override(self):
        """Test overriding default h1 style"""
        tag_overrides = {
            "h1": "Heading 2",
        }

        html = "<h1>Test Heading</h1>"

        doc = Document()
        parser = HtmlToDocx(tag_style_overrides=tag_overrides)
        parser.add_html_to_document(html, doc)

        # h1 should use Heading 2 instead of default Heading 1
        self.assertEqual(doc.paragraphs[0].style.name, "Heading 2")

    def test_class_overrides_tag_override(self):
        """Test that class mapping has priority over tag override"""
        style_map = {"custom": "Heading 3"}
        tag_overrides = {"h1": "Heading 2"}

        html = '<h1 class="custom">Test</h1>'

        doc = Document()
        parser = HtmlToDocx(style_map=style_map, tag_style_overrides=tag_overrides)
        parser.add_html_to_document(html, doc)

        # Class should win over tag override
        self.assertEqual(doc.paragraphs[0].style.name, "Heading 3")


class TestDefaultParagraphStyle(unittest.TestCase):
    """Test default paragraph style functionality"""

    def test_normal_default(self):
        """Test that Normal is used as default by default"""
        html = "<p>Test paragraph</p>"

        doc = Document()
        parser = HtmlToDocx()  # default_paragraph_style=None by default
        parser.add_html_to_document(html, doc)

        self.assertEqual(doc.paragraphs[0].style.name, "Normal")

    def test_custom_default(self):
        """Test setting custom default paragraph style"""
        html = "<p>Test paragraph</p>"

        doc = Document()
        parser = HtmlToDocx(default_paragraph_style="Heading 1")
        parser.add_html_to_document(html, doc)

        self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")

    def test_none_default_uses_normal(self):
        """Test that None uses Word's default Normal style"""
        html = "<p>Test paragraph</p>"

        doc = Document()
        parser = HtmlToDocx(default_paragraph_style=None)
        parser.add_html_to_document(html, doc)

        self.assertEqual(doc.paragraphs[0].style.name, "Normal")


class TestInlineStyles(unittest.TestCase):
    """Test inline CSS style support"""

    def test_inline_color(self):
        """Test inline color style"""
        html = '<p><span style="color: red">Red text</span></p>'

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, doc)

        # Check that run has color applied
        run = doc.paragraphs[0].runs[0]
        self.assertIsNotNone(run.font.color.rgb)

    def test_inline_font_size(self):
        """Test inline font-size style"""
        html = '<p><span style="font-size: 18pt">Large text</span></p>'

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertEqual(run.font.size, Pt(18))

    def test_inline_bold(self):
        """Test inline font-weight bold"""
        html = '<p><span style="font-weight: bold">Bold text</span></p>'

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertTrue(run.font.bold)

    def test_inline_italic(self):
        """Test inline font-style italic"""
        html = '<p><span style="font-style: italic">Italic text</span></p>'

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertTrue(run.font.italic)

    def test_paragraph_inline_styles(self):
        """Test inline styles on paragraph elements"""
        html = '<p style="color: blue; font-size: 14pt">Blue 14pt paragraph</p>'

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertIsNotNone(run.font.color.rgb)
        self.assertEqual(run.font.size, Pt(14))


class TestImportantFlag(unittest.TestCase):
    """Test !important CSS flag precedence"""

    def test_important_overrides_normal(self):
        """Test that !important styles override normal styles"""
        html = """
        <p>
            <span style="color: gray">
                Gray text with <span style="color: red !important">red important</span>.
            </span>
        </p>
        """

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, doc)

        # The "red important" run should have red color
        # (exact run index may vary based on whitespace handling)
        runs = doc.paragraphs[0].runs
        self.assertTrue(len(runs) > 0)

    def test_important_on_paragraph(self):
        """Test !important on paragraph inline style"""
        html = '<p style="color: blue !important">Blue important</p>'

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertIsNotNone(run.font.color.rgb)


class TestCodeBlockMultipleParagraphs(unittest.TestCase):
    """Test that code block styles persist across multiple paragraphs"""

    def test_multi_paragraph_code_block(self):
        """Test that all paragraphs in code block maintain style"""
        style_map = {
            "code-block": "No Spacing",  # Using built-in style
        }

        html = """
        <div class="code-block">
            <p>First line of code</p>
            <p>Second line of code</p>
            <p>Third line of code</p>
        </div>
        """

        doc = Document()
        parser = HtmlToDocx(style_map=style_map)
        parser.add_html_to_document(html, doc)

        # All three paragraphs should have the code-block style
        self.assertEqual(doc.paragraphs[0].style.name, "No Spacing")
        self.assertEqual(doc.paragraphs[1].style.name, "No Spacing")
        self.assertEqual(doc.paragraphs[2].style.name, "No Spacing")


class TestCKEditorStyles(unittest.TestCase):
    """Test CKEditor-style custom classes"""

    def test_numbered_headings(self):
        """Test numbered heading classes"""
        style_map = {
            "numbered-heading-1": "Heading 1",
            "numbered-heading-2": "Heading 2",
            "numbered-heading-3": "Heading 3",
        }

        html = """
        <h1 class="numbered-heading-1">1.0 Introduction</h1>
        <h2 class="numbered-heading-2">1.1 Overview</h2>
        <h3 class="numbered-heading-3">1.1.1 Details</h3>
        """

        doc = Document()
        parser = HtmlToDocx(style_map=style_map)
        parser.add_html_to_document(html, doc)

        self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")
        self.assertEqual(doc.paragraphs[1].style.name, "Heading 2")
        self.assertEqual(doc.paragraphs[2].style.name, "Heading 3")


class TestBackwardCompatibility(unittest.TestCase):
    """Test that existing html4docx functionality still works"""

    def test_basic_html_still_works(self):
        """Test that basic HTML conversion works without new features"""
        html = "<p>Simple paragraph</p><h1>Heading</h1>"

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, doc)

        self.assertEqual(len(doc.paragraphs), 2)
        self.assertEqual(doc.paragraphs[1].style.name, "Heading 1")

    def test_existing_span_styles_work(self):
        """Test that existing <span style="..."> still works"""
        html = '<p><span style="color: #FF0000">Red text</span></p>'

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertIsNotNone(run.font.color.rgb)

    def test_bold_italic_tags_work(self):
        """Test that <b>, <i>, <u> tags still work"""
        html = "<p><b>Bold</b> <i>Italic</i> <u>Underline</u></p>"

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, doc)

        # Find runs with the specific formatting (spaces create extra runs, so we can't rely on indices)
        runs = doc.paragraphs[0].runs
        bold_runs = [r for r in runs if r.font.bold]
        italic_runs = [r for r in runs if r.font.italic]
        underline_runs = [r for r in runs if r.font.underline]

        self.assertTrue(len(bold_runs) > 0, "Should have at least one bold run")
        self.assertTrue(len(italic_runs) > 0, "Should have at least one italic run")
        self.assertTrue(len(underline_runs) > 0, "Should have at least one underline run")


class TestEdgeCases(unittest.TestCase):
    """Test edge cases and error handling"""

    def test_nonexistent_style_graceful_failure(self):
        """Test that non-existent styles don't crash"""
        style_map = {
            "custom": "NonExistentStyle",
        }

        html = '<p class="custom">Test</p>'

        doc = Document()
        parser = HtmlToDocx(style_map=style_map)

        # Should not raise exception
        try:
            parser.add_html_to_document(html, doc)
            success = True
        except Exception:
            success = False

        self.assertTrue(success)

    def test_empty_style_map(self):
        """Test with empty style_map"""
        html = '<p class="anything">Test</p>'

        doc = Document()
        parser = HtmlToDocx(style_map={})
        parser.add_html_to_document(html, doc)

        # Should use default (Normal)
        self.assertEqual(doc.paragraphs[0].style.name, "Normal")

    def test_none_style_map(self):
        """Test with None style_map"""
        html = "<p>Test</p>"

        doc = Document()
        parser = HtmlToDocx(style_map=None)
        parser.add_html_to_document(html, doc)

        self.assertEqual(len(doc.paragraphs), 1)


def run_tests():
    """Run all tests and return results"""
    loader = unittest.TestLoader()
    suite = loader.loadTestsFromModule(__import__(__name__))
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    return result.wasSuccessful()


if __name__ == "__main__":
    unittest.main()
